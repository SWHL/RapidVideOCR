# !/usr/bin/env python
# -*- encoding: utf-8 -*-
# @Author: SWHL
# @Contact: liekkaskono@163.com
import copy
import datetime
import difflib
from pathlib import Path
from io import BytesIO

import cv2
import numpy as np
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def compute_centroid(dt_box):
    x = (np.max(dt_box[:, 0]) + np.min(dt_box[:, 0])) / 2
    y = (np.max(dt_box[:, 1]) + np.min(dt_box[:, 1])) / 2
    return (x, y)


def compute_height(dt_box):
    return np.max(dt_box[:, 1]) - np.min(dt_box[:, 1])


def is_two_lines(dt_boxes):
    threshold = max([compute_height(dt_box) for dt_box in dt_boxes])
    centroid_list = [compute_centroid(dt_box) for dt_box in dt_boxes]
    first_point = centroid_list[0]
    second_point = centroid_list[-1]
    distance = np.sqrt(np.abs(first_point[0] - second_point[0]) ** 2
                       + np.abs(first_point[1] - second_point[1]) ** 2)
    if distance > threshold:
        return False
    else:
        return True


def string_similar(s1, s2):
    return difflib.SequenceMatcher(None, s1, s2).ratio()


def get_frame_from_time(time_str, fps):
    if time_str:
        frame_index = get_frame_index(time_str, fps)
    else:
        frame_index = 0
    return frame_index


def get_frame_index(time_str: str, fps: float):
    # convert time string to frame index
    t = time_str.split(':')
    t = list(map(float, t))
    if len(t) == 3:
        td = datetime.timedelta(hours=t[0], minutes=t[1], seconds=t[2])
    elif len(t) == 2:
        td = datetime.timedelta(minutes=t[0], seconds=t[1])
    else:
        raise ValueError(
            f'Time data "{time_str}" does not match format "%H:%M:%S"')
    index = int(td.total_seconds() * fps)
    return index


def get_srt_timestamp(frame_index: int, fps: float):
    # convert frame index into SRT timestamp
    td = datetime.timedelta(seconds=frame_index / fps)
    ms = td.microseconds // 1000
    m, s = divmod(td.seconds, 60)
    h, m = divmod(m, 60)
    return f'{h:02d}:{m:02d}:{s:02d},{ms:03d}', td.seconds * 1000 + ms


def is_similar_batch(img_a, img_batch, threshold=0.000):
    img_a_tmp = copy.deepcopy(img_a)
    img_batch_tmp = copy.deepcopy(img_batch)

    img_a_tmp /= 255
    img_batch_tmp /= 255

    difference = (img_a_tmp - img_batch_tmp) ** 2
    difference = difference.reshape(img_batch_tmp.shape[0], -1)

    error = np.sum(difference, axis=1) / img_a_tmp.size
    return error < threshold


def rgb_to_grey(img):
    if img.ndim == 3:
        img = img[np.newaxis, :, :, :]
    return img[..., 0] * 0.114 + img[..., 1] * 0.587 + img[..., 2] * 0.299


def binary_img(img, binary_threshold=243):
    _, img = cv2.threshold(img, binary_threshold, 255, cv2.THRESH_BINARY)
    return img


def vis_binary(img):
    img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    img2 = copy.deepcopy(img)

    def update_theta(x): pass

    window_name = 'image'
    tracker_name = 'threshold'

    cv2.namedWindow(window_name)
    cv2.createTrackbar(tracker_name, window_name, 0, 255, update_theta)
    cv2.setTrackbarPos(trackbarname=tracker_name,
                       winname=window_name,
                       pos=127)

    while (True):
        cv2.imshow(window_name, img)

        threshold = cv2.getTrackbarPos(tracker_name, window_name)
        _, img = cv2.threshold(img2, threshold, 255, cv2.THRESH_BINARY)

        if cv2.waitKey(1) == ord('q'):
            break
    cv2.destroyAllWindows()
    return threshold


def dilate_img(img):
    kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
    img = cv2.dilate(img, kernel, iterations=1)
    return img


def remove_bg(img, is_dilate=True, binary_thr=243):
    img = rgb_to_grey(img).squeeze()
    if is_dilate:
        img = dilate_img(binary_img(img, binary_thr))
    img = img[np.newaxis, :, :]
    return img


def remove_batch_bg(img_batch, is_dilate=True, binary_thr=243):
    img_batch = rgb_to_grey(img_batch)
    new_img_batch = []
    for img_one in img_batch:
        if is_dilate:
            img_one = dilate_img(binary_img(img_one, binary_thr))

        new_img_batch.append(img_one)
    img_batch = np.array(new_img_batch)
    return img_batch


def save_srt(video_path, extract_result):
    final_result = []
    for i, start_index, end_index, text in extract_result:
        final_result.append(
            f'{i}\n{start_index} --> {end_index}\n{text}\n')

    save_output_dir = Path(video_path).parent
    video_path_stem = Path(video_path).stem
    save_full_path = save_output_dir / f'{video_path_stem}.srt'

    with open(save_full_path, 'w', encoding='utf-8') as f:
        for value in final_result:
            f.write(value + '\n')
    print(f'The srt has been saved in the {save_full_path}.')


def save_txt(video_path, extract_result):
    final_result = []
    for _, _, _, text in extract_result:
        final_result.append(f'{text}\n')

    save_output_dir = Path(video_path).parent
    video_path_stem = Path(video_path).stem
    save_full_path = save_output_dir / f'{video_path_stem}.txt'

    with open(save_full_path, 'w', encoding='utf-8') as f:
        for value in final_result:
            f.write(value + '\n')
    print(f'The txt has been saved in the {save_full_path}.')


def save_docx(video_path, extract_result, vr):
    """
        将带有图像和对应的文本保存到word中
        图像在上，文字在下，居中排列
    """
    doc = Document()
    for k, _, _, text in extract_result:
        raw_im = cv2.cvtColor(vr[k].asnumpy(), cv2.COLOR_BGR2RGB)
        im = cv2.imencode('.jpg', raw_im)[1]
        im_bytes = BytesIO(im.tobytes())

        img_para = doc.add_paragraph()
        img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        img_run = img_para.add_run('')
        img_run.add_picture(im_bytes, width=Cm(13.93))

        text_para = doc.add_paragraph()
        text_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        text_para.add_run(text)

    save_output_dir = Path(video_path).parent
    video_path_stem = Path(video_path).stem
    save_full_path = save_output_dir / f'{video_path_stem}.docx'

    doc.save(str(save_full_path))
    print(f'The docx has been saved in the {save_full_path}.')


def debug_vis_box(i, dt_boxes, one_frame):
    for box in dt_boxes:
        box = np.array(box).astype(np.int32).reshape(-1, 2)
        cv2.polylines(one_frame, [box], True,
                      color=(255, 255, 0), thickness=2)
    cv2.imwrite(f'temp/{i}.jpg', one_frame)

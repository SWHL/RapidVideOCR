
# !/usr/bin/env python
# -*- encoding: utf-8 -*-
# @Author: SWHL
# @Contact: liekkaskono@163.com
from pathlib import Path

import cv2
import numpy as np
from decord import VideoReader, cpu
from docx import Document
from docx.api import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from io import BytesIO
from tqdm import tqdm

from .utils import (get_frame_from_time, get_srt_timestamp, is_similar_batch,
                    remove_batch_bg, remove_bg, string_similar)


class Video(object):
    def __init__(self,
                 video_path,
                 ocr_system,
                 batch_size=None,
                 subtitle_height=45,
                 error_num=0.1,
                 output_format='srt'):
        self.video_path = video_path
        self.ocr_system = ocr_system
        self.error_num = error_num
        self.output_format = output_format

        print(f'Loading {self.video_path}')
        self.vr = VideoReader(self.video_path, ctx=cpu(0))
        self.num_frames = len(self.vr)
        self.fps = int(self.vr.get_avg_fps())

        self.height = int(self.vr[0].shape[0])
        self.subtitle_height = subtitle_height
        self.crop_h = self.height - self.subtitle_height

        if batch_size is None:
            self.batch_size = self.fps * 2
        else:
            self.batch_size = batch_size

    def record_info(self, slow, duplicate_frame, slow_frame):
        # Record
        if slow in self.key_point_dict:
            self.key_point_dict[slow].extend(duplicate_frame)
        else:
            self.key_point_dict[slow] = duplicate_frame

        if slow not in self.key_frames.keys():
            self.key_frames[slow] = slow_frame

    def get_key_frame(self):
        self.key_point_dict = {}
        self.key_frames = {}
        with tqdm(total=self.ocr_end, desc='Get the key frame') as pbar:
            # Use two fast and slow pointers to filter duplicate frame.
            if self.batch_size > self.ocr_end:
                self.batch_size = self.ocr_end - 1

            slow, fast = 0, 1
            while fast + self.batch_size <= self.ocr_end:
                slow_frame = self.vr[slow].asnumpy()[self.crop_h:, :, :]

                batch_list = list(range(fast, fast+self.batch_size))
                fast_frames = self.vr.get_batch(batch_list).asnumpy()
                fast_frames = fast_frames[:, self.crop_h:, :, :]

                # Remove the background of the frame.
                slow_frame = remove_bg(slow_frame)
                fast_frames = remove_batch_bg(fast_frames)

                # Compare the similarity between the frames.
                compare_result = is_similar_batch(slow_frame,
                                                  fast_frames,
                                                  threshold=self.error_num)
                batch_array = np.array(batch_list)
                not_similar_index = batch_array[np.logical_not(compare_result)]

                duplicate_frame = []
                if not_similar_index.size == 0:
                    # All are similar with the slow frame.
                    duplicate_frame = batch_list

                    pbar.update(self.batch_size)
                    fast += self.batch_size

                    self.record_info(slow, duplicate_frame, slow_frame)
                else:
                    # Exist the non similar frame.
                    index = not_similar_index[0] - slow
                    duplicate_frame = batch_list[:index]

                    # record
                    self.record_info(slow, duplicate_frame, slow_frame)

                    slow = not_similar_index[0]
                    fast = slow + 1
                    pbar.update(slow - pbar.n + 1)

                # Take care the left frames, which can't up to the batch_size.
                if fast != self.ocr_end \
                        and fast + self.batch_size > self.ocr_end:
                    self.batch_size = self.ocr_end - fast

    def run_ocr(self, time_start, time_end):
        self.ocr_start = get_frame_from_time(time_start, self.fps)

        if time_end == '-1':
            self.ocr_end = self.num_frames - 1
        else:
            self.ocr_end = get_frame_from_time(time_end, self.fps)

        if self.ocr_end < self.ocr_start:
            raise ValueError('time_start is later than time_end')

        self.get_key_frame()

        # Extract the filtered frames content.
        self.pred_frames = []
        key_index_frames = list(self.key_point_dict.keys())
        frames = np.stack(list(self.key_frames.values()), axis=0)

        for key, frame in tqdm(list(zip(key_index_frames, frames)),
                               desc='Extract content'):
            frame = cv2.copyMakeBorder(frame.squeeze(),
                                       self.subtitle_height * 2,
                                       self.subtitle_height * 2,
                                       0, 0,
                                       cv2.BORDER_CONSTANT,
                                       value=(0, 0))
            frame = cv2.cvtColor(frame.astype(np.uint8), cv2.COLOR_GRAY2BGR)

            _, rec_res = self.ocr_system(frame)

            if rec_res is None or len(rec_res) <= 0:
                del self.key_point_dict[key]
            else:
                text, confidence = list(zip(*rec_res))
                text = '\n'.join(text)
                confidence = sum(confidence) / len(confidence)
                self.pred_frames.append((text, confidence))

    def get_subtitles(self):
        slow, fast = 0, 1
        n = len(self.pred_frames)
        key_list = list(self.key_point_dict.keys())
        invalid_keys = []
        while fast < n:
            if string_similar(self.pred_frames[slow][0],
                              self.pred_frames[fast][0]) > 0.6:
                # 相似→合并两个list
                self.key_point_dict[key_list[slow]].extend(
                    self.key_point_dict[key_list[fast]])
                self.key_point_dict[key_list[slow]].sort()
                invalid_keys.append(fast)
            else:
                # 不相似
                slow = fast
            fast += 1

        [self.key_point_dict.pop(key_list[i]) for i in invalid_keys]
        self.pred_frames = [v for i, v in enumerate(self.pred_frames)
                            if i not in invalid_keys]

        self.extract_result = []
        for i, (k, v) in enumerate(self.key_point_dict.items()):
            start_index = get_srt_timestamp(v[0], self.fps)
            end_index = get_srt_timestamp(v[-1], self.fps)
            self.extract_result.append([k, start_index,
                                        end_index, self.pred_frames[i][0]])
        self.save_output()
        return self.extract_result

    def save_output(self):
        if self.output_format == 'srt':
            self._save_srt()
        elif self.output_format == 'txt':
            self._save_txt()
        elif self.output_format == 'docx':
            self._save_docx()
        elif self.output_format == 'all':
            self._save_srt()
            self._save_txt()
            self._save_docx()
        else:
            raise ValueError(f'The {self.output_format} is not supported!')

    def _save_srt(self):
        final_result = []
        for i, start_index, end_index, text in self.extract_result:
            final_result.append(
                f'{i}\n{start_index} --> {end_index}\n{text}\n')

        save_output_dir = Path(self.video_path).parent
        video_path_stem = Path(self.video_path).stem
        save_full_path = save_output_dir / f'{video_path_stem}.srt'

        with open(save_full_path, 'w', encoding='utf-8') as f:
            for value in final_result:
                f.write(value + '\n')
        print(f'The srt has been saved in the {save_full_path}.')

    def _save_txt(self):
        final_result = []
        for _, _, _, text in self.extract_result:
            final_result.append(f'{text}\n')

        save_output_dir = Path(self.video_path).parent
        video_path_stem = Path(self.video_path).stem
        save_full_path = save_output_dir / f'{video_path_stem}.txt'

        with open(save_full_path, 'w', encoding='utf-8') as f:
            for value in final_result:
                f.write(value + '\n')
        print(f'The txt has been saved in the {save_full_path}.')

    def _save_docx(self):
        """
            将带有图像和对应的文本保存到word中
            图像在上，文字在下，居中排列
        """
        doc = Document()
        for k, _, _, text in self.extract_result:
            raw_im = cv2.cvtColor(self.vr[k].asnumpy(), cv2.COLOR_BGR2RGB)
            im = cv2.imencode('.jpg', raw_im)[1]
            im_bytes = BytesIO(im.tobytes())

            img_para = doc.add_paragraph()
            img_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            img_run = img_para.add_run('')
            img_run.add_picture(im_bytes, width=Cm(13.93))

            text_para = doc.add_paragraph()
            text_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            text_para.add_run(text)

        save_output_dir = Path(self.video_path).parent
        video_path_stem = Path(self.video_path).stem
        save_full_path = save_output_dir / f'{video_path_stem}.docx'

        doc.save(str(save_full_path))
        print(f'The docx has been saved in the {save_full_path}.')
